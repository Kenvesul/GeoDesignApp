"""
exporters/report_docx.py
========================
Editable Microsoft Word (.docx) calculation sheet for slope stability analysis.

Mirrors the content of report_pdf.py but produces a fully editable .docx that
engineers can annotate, stamp, and incorporate into project documentation.

Sections
--------
1. Project header block
2. Input parameters table
3. EC7 DA1 verification table (Combination 1 & 2)
4. Verdict stamp (PASS / FAIL)
5. Slice summary table (first 12 slices)
6. Embedded cross-section figure (PNG via plot_slope)

Standards & references
----------------------
- EC7 EN 1997-1:2004  (partial factors, DA1 combinations)
- python-docx 1.x API (https://python-docx.readthedocs.io/)

Dependencies (exporters only — not core/)
-----------------------------------------
- python-docx  >= 1.0
- matplotlib   (via plot_slope.plot_slope_stability for the embedded figure)
"""

from __future__ import annotations

import io
import math
import datetime
import tempfile
import os
from typing import TYPE_CHECKING

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

if TYPE_CHECKING:
    from models.soil import Soil
    from models.geometry import SlopeGeometry
    from core.search import SearchResult
    from core.factors_of_safety import VerificationResult
    from core.slicer import Slice

# ── Colour palette (matches report_pdf.py) ──────────────────────────────────
_C_HEADER_BG   = "1A3A5C"   # dark navy
_C_HEADER_FG   = "FFFFFF"   # white
_C_SUBHDR_BG   = "D6E8F7"   # light blue
_C_PASS        = "1D6A2B"   # dark green
_C_FAIL        = "B91C1C"   # dark red
_C_PASS_BG     = "D4EFDF"   # light green
_C_FAIL_BG     = "FADBD8"   # light red
_C_BORDER      = "7F95AD"   # medium blue-grey


# ── Low-level XML helpers ────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_colour: str) -> None:
    """Set solid background fill on a table cell using raw OOXML.

    python-docx does not expose cell shading through a high-level API;
    direct XML manipulation is the canonical approach (python-docx FAQ).

    Parameters
    ----------
    cell       : docx TableCell
    hex_colour : 6-character hex string, e.g. '1A3A5C'  (no leading #)
    """
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_colour)
    tcPr.append(shd)


def _set_cell_borders(cell, hex_colour: str = _C_BORDER) -> None:
    """Apply thin borders to all four sides of a cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tbl_brd = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")        # ½ pt
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), hex_colour)
        tbl_brd.append(el)
    tcPr.append(tbl_brd)


def _cell_para(cell,
               text: str,
               bold: bool = False,
               italic: bool = False,
               fg: str | None = None,
               align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
               size_pt: float = 9.0) -> None:
    """Write text into a cell's first paragraph with optional formatting."""
    para = cell.paragraphs[0]
    para.alignment = align
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after  = Pt(1)
    run  = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size_pt)
    if fg:
        run.font.color.rgb = RGBColor.from_string(fg)


# ── Document-level helpers ───────────────────────────────────────────────────

def _apply_normal_style(doc: Document) -> None:
    """Override Normal style: Arial 10 pt, tight spacing."""
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after  = Pt(4)


def _section_heading(doc: Document, text: str) -> None:
    """Add a styled section heading paragraph (no built-in Heading styles needed)."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after  = Pt(4)
    run  = para.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(_C_HEADER_BG)
    # Add bottom border to mimic a ruled heading
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), _C_HEADER_BG)
    pBdr.append(bot)
    pPr.append(pBdr)


# ── Table builders ───────────────────────────────────────────────────────────

def _header_table(doc: Document,
                  project: str,
                  job_ref: str,
                  calc_by: str,
                  checked_by: str,
                  date: str) -> None:
    """
    Two-column metadata block at the top of the document.

    Layout (3 rows × 2 cols):
        | DesignApp — Slope Stability   | Project: <name>     |
        | EC7 EN 1997-1:2004 — DA1      | Job Ref: <ref>      |
        | Date: <date>                  | Calc: <x>  Chk: <y> |
    """
    tbl = doc.add_table(rows=3, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    col_w = [Cm(9.5), Cm(9.5)]

    # Row 0: title / project
    row0 = tbl.rows[0]
    _set_cell_bg(row0.cells[0], _C_HEADER_BG)
    _set_cell_bg(row0.cells[1], _C_HEADER_BG)
    _cell_para(row0.cells[0],
               "DesignApp — Slope Stability Calculation",
               bold=True, fg=_C_HEADER_FG, size_pt=11)
    _cell_para(row0.cells[1],
               f"Project: {project}",
               bold=True, fg=_C_HEADER_FG, align=WD_ALIGN_PARAGRAPH.RIGHT)

    # Row 1: standard / job ref
    row1 = tbl.rows[1]
    _set_cell_bg(row1.cells[0], _C_SUBHDR_BG)
    _set_cell_bg(row1.cells[1], _C_SUBHDR_BG)
    _cell_para(row1.cells[0], "Eurocode 7  EN 1997-1:2004  ·  Design Approach 1", italic=True)
    _cell_para(row1.cells[1], f"Job Ref: {job_ref}", align=WD_ALIGN_PARAGRAPH.RIGHT)

    # Row 2: date / calc-checked
    row2 = tbl.rows[2]
    _set_cell_bg(row2.cells[0], _C_SUBHDR_BG)
    _set_cell_bg(row2.cells[1], _C_SUBHDR_BG)
    _cell_para(row2.cells[0], f"Date: {date}")
    _cell_para(row2.cells[1],
               f"Calc: {calc_by}   Chk: {checked_by}",
               align=WD_ALIGN_PARAGRAPH.RIGHT)

    # Set column widths
    for i, row in enumerate(tbl.rows):
        for j, cell in enumerate(row.cells):
            cell.width = col_w[j]
            _set_cell_borders(cell)

    doc.add_paragraph()  # spacer


def _input_table(doc: Document,
                 soil: "Soil",
                 slope: "SlopeGeometry",
                 circle,
                 method: str,
                 ru: float) -> None:
    """
    Two-column table: Parameter | Value.

    Soil properties from: Craig (2004) Ch. 9, EC7 §2.4.6
    Circle notation:      cx, cy = centre coords (m), R = radius (m)
    """
    rows_data = [
        # (label, value)
        ("Soil name",           soil.name),
        ("Unit weight  γ",      f"{soil.gamma:.1f}  kN/m³"),
        ("Friction angle  φ'k", f"{soil.phi_k:.1f}°"),
        ("Cohesion  c'k",       f"{soil.c_k:.1f}  kPa"),
        ("Slope points",        "  →  ".join(f"({x:.1f}, {y:.1f})" for x, y in slope.points)),
        ("Critical circle cx",  f"{circle.cx:.3f}  m"),
        ("Critical circle cy",  f"{circle.cy:.3f}  m"),
        ("Critical circle R",   f"{circle.r:.3f}  m"),
        ("Analysis method",     method),
        ("Pore pressure ratio ru", f"{ru:.3f}"),
    ]

    headers = ["Parameter", "Value"]
    tbl = doc.add_table(rows=1 + len(rows_data), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr = tbl.rows[0].cells
    for j, h in enumerate(headers):
        _set_cell_bg(hdr[j], _C_SUBHDR_BG)
        _cell_para(hdr[j], h, bold=True)
        _set_cell_borders(hdr[j])

    # Data rows
    for i, (label, value) in enumerate(rows_data):
        cells = tbl.rows[i + 1].cells
        _cell_para(cells[0], label)
        _cell_para(cells[1], value)
        for cell in cells:
            _set_cell_borders(cell)

    # Column widths
    for row in tbl.rows:
        row.cells[0].width = Cm(7)
        row.cells[1].width = Cm(12)

    doc.add_paragraph()


def _da1_table(doc: Document, verification: "VerificationResult") -> None:
    """
    EC7 DA1 verification table.

    Columns: Combination | γφ | φ'd (°) | c'd (kPa) | FoS_d | Result
    Ref: EC7 EN 1997-1:2004 §2.4.7.3, Table A.3/A.4
    """
    headers = ["Combination", "γφ", "φ'_d (°)", "c'_d (kPa)", "FoS_d", "Result"]
    combos  = [verification.comb1, verification.comb2]
    labels  = ["C1 (A1+M1+R1)", "C2 (A2+M2+R1)"]

    tbl = doc.add_table(rows=1 + len(combos), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr = tbl.rows[0].cells
    for j, h in enumerate(headers):
        _set_cell_bg(hdr[j], _C_SUBHDR_BG)
        _cell_para(hdr[j], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_borders(hdr[j])

    # Data rows
    for i, (combo, label) in enumerate(zip(combos, labels)):
        cells = tbl.rows[i + 1].cells
        pass_fail = combo.passes

        bg  = _C_PASS_BG if pass_fail else _C_FAIL_BG
        fg  = _C_PASS    if pass_fail else _C_FAIL
        txt = "PASS ✓" if pass_fail else "FAIL ✗"

        row_data = [
            label,
            f"{combo.gamma_phi:.2f}",
            f"{combo.phi_d:.1f}",
            f"{combo.c_d:.2f}",
            f"{combo.fos_d:.3f}",
            txt,
        ]
        for j, val in enumerate(row_data):
            _set_cell_borders(cells[j])
            if j == 5:
                _set_cell_bg(cells[j], bg)
                _cell_para(cells[j], val, bold=True, fg=fg,
                           align=WD_ALIGN_PARAGRAPH.CENTER)
            else:
                _cell_para(cells[j], val, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Column widths
    col_widths = [Cm(4.5), Cm(1.8), Cm(2.2), Cm(2.8), Cm(2.2), Cm(2.5)]
    for row in tbl.rows:
        for j, w in enumerate(col_widths):
            row.cells[j].width = w

    doc.add_paragraph()


def _verdict_paragraph(doc: Document, verification: "VerificationResult") -> None:
    """
    Bold PASS / FAIL summary sentence with governing FoS.

    EC7 requirement: both DA1 combinations must pass (§2.4.7.3.4(2)P).
    """
    overall   = verification.passes
    gov_combo = verification.comb2 if verification.comb2.fos_d < verification.comb1.fos_d else verification.comb1
    gov_label = "C2" if verification.comb2.fos_d < verification.comb1.fos_d else "C1"

    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(8)

    label_run = para.add_run("EC7 DA1 Verdict: ")
    label_run.bold      = True
    label_run.font.size = Pt(11)

    verdict_run = para.add_run(
        "SATISFACTORY ✓" if overall else "UNSATISFACTORY ✗"
    )
    verdict_run.bold      = True
    verdict_run.font.size = Pt(11)
    verdict_run.font.color.rgb = RGBColor.from_string(
        _C_PASS if overall else _C_FAIL
    )

    suffix_run = para.add_run(
        f"    (Governing: {gov_label}, FoS_d = {gov_combo.fos_d:.3f})"
    )
    suffix_run.font.size = Pt(10)


def _slice_table(doc: Document, slices: list["Slice"], max_rows: int = 12) -> None:
    """
    Tabulate the first *max_rows* vertical slices.

    Columns: #, x_mid (m), b (m), α (°), W (kN/m), sin α, cos α
    Source: Craig (2004) §9.3 — Bishop's Simplified Method slice geometry.
    """
    subset  = slices[:max_rows]
    headers = ["#", "x (m)", "b (m)", "α (°)", "W (kN/m)", "sin α", "cos α"]

    tbl = doc.add_table(rows=1 + len(subset), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr = tbl.rows[0].cells
    for j, h in enumerate(headers):
        _set_cell_bg(hdr[j], _C_SUBHDR_BG)
        _cell_para(hdr[j], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_borders(hdr[j])

    # Data rows
    for i, s in enumerate(subset):
        cells  = tbl.rows[i + 1].cells
        values = [
            str(i + 1),
            f"{s.x:.3f}",
            f"{s.b:.3f}",
            f"{math.degrees(s.alpha):.2f}",
            f"{s.weight:.2f}",
            f"{math.sin(s.alpha):.4f}",
            f"{math.cos(s.alpha):.4f}",
        ]
        for j, val in enumerate(values):
            _cell_para(cells[j], val, align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_borders(cells[j])

    # Uniform column widths
    col_w = Cm(19 / len(headers))
    for row in tbl.rows:
        for cell in row.cells:
            cell.width = col_w

    doc.add_paragraph()


# ── Main public function ─────────────────────────────────────────────────────

def generate_slope_report_docx(
    filepath:       str,
    soil:           "Soil",
    slope:          "SlopeGeometry",
    search_result:  "SearchResult",
    verification:   "VerificationResult",
    slices:         list["Slice"],
    ru:             float = 0.0,
    project:        str   = "Untitled Project",
    job_ref:        str   = "",
    calc_by:        str   = "",
    checked_by:     str   = "",
    date:           str   = "",
) -> None:
    """
    Generate an editable Word (.docx) slope stability calculation sheet.

    The document mirrors the PDF produced by report_pdf.generate_slope_report
    but is fully editable in Microsoft Word / LibreOffice.

    Parameters
    ----------
    filepath       : Output path, e.g. '/tmp/slope_calc.docx'
    soil           : Soil model (Soil dataclass, models/soil.py)
    slope          : Ground surface geometry (SlopeGeometry, models/geometry.py)
    search_result  : Grid/golden-section search result (SearchResult, core/search.py)
    verification   : EC7 DA1 verification result (VerificationResult,
                     core/factors_of_safety.py)
    slices         : List of Slice objects (core/slicer.py)
    ru             : Pore pressure ratio (dimensionless)
    project        : Project name for header
    job_ref        : Job reference number
    calc_by        : Calculator's initials / name
    checked_by     : Checker's initials / name
    date           : Date string (defaults to today if empty)

    Returns
    -------
    None  —  file is written to *filepath*.

    References
    ----------
    EC7 §2.4.7.3   Design Approach 1
    EC7 Table A.3  Partial factors on actions (A1, A2)
    EC7 Table A.4  Partial factors on material properties (M1, M2)
    Bishop (1955)  Simplified method of analysis
    Craig (2004)   Soil Mechanics, Ch. 9
    """
    if not date:
        date = datetime.date.today().isoformat()

    circle = search_result.critical_circle
    method = search_result.best_fos_result.method  # e.g. "Bishop Simplified"

    doc = Document()
    _apply_normal_style(doc)

    # ── A4 page setup ────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width   = Cm(21.0)
    section.page_height  = Cm(29.7)
    section.left_margin  = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin   = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # ── 1. Header block ──────────────────────────────────────────────────────
    _header_table(doc, project, job_ref, calc_by, checked_by, date)

    # ── 2. Input parameters ──────────────────────────────────────────────────
    _section_heading(doc, "1.  Input Parameters")
    _input_table(doc, soil, slope, circle, method, ru)

    # ── 3. EC7 DA1 verification ──────────────────────────────────────────────
    _section_heading(doc, "2.  EC7 DA1 Verification  (EN 1997-1:2004 §2.4.7.3)")
    _da1_table(doc, verification)
    _verdict_paragraph(doc, verification)

    # ── 4. Slice summary ─────────────────────────────────────────────────────
    _section_heading(doc, "3.  Slice Summary  (Bishop 1955 — first 12 slices)")
    _slice_table(doc, slices)

    # ── 5. Cross-section figure ──────────────────────────────────────────────
    _section_heading(doc, "4.  Slope Cross-Section")
    try:
        from exporters.plot_slope import plot_slope_stability
        import matplotlib
        matplotlib.use("Agg")

        fig = plot_slope_stability(slope, search_result, title="Critical Slip Circle", ru=ru)

        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
            tmp_path = tmp.name
        fig.savefig(tmp_path, dpi=150, bbox_inches="tight")
        import matplotlib.pyplot as plt
        plt.close(fig)

        doc.add_picture(tmp_path, width=Cm(17.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        os.unlink(tmp_path)

    except Exception as exc:  # graceful degradation — figure is optional
        note = doc.add_paragraph(f"[Figure not available: {exc}]")
        note.runs[0].italic = True

    # ── Footer note ──────────────────────────────────────────────────────────
    doc.add_paragraph()
    footer_para = doc.add_paragraph(
        "Generated by DesignApp · Eurocode 7 EN 1997-1:2004 · "
        "Results are characteristic unless noted.  "
        "Verify all inputs before using for design."
    )
    footer_para.runs[0].font.size = Pt(7)
    footer_para.runs[0].italic    = True
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(filepath)


# ============================================================
#  Foundation bearing capacity DOCX report  (B-07)
# ============================================================

def generate_foundation_report_docx(
    filepath    : str,
    analysis    : dict,
    project     : str = "Untitled Project",
    job_ref     : str = "—",
    calc_by     : str = "DesignApp",
    checked_by  : str = "—",
) -> None:
    """
    Editable Word (.docx) calculation sheet for EC7 DA1 foundation analysis.

    Mirrors generate_foundation_report() in report_pdf.py but produces
    a fully-editable .docx suitable for engineer annotation and issue.

    Sections:
        1. Project header block
        2. Input parameters table
        3. EC7 DA1 bearing capacity (ULS) table
        4. Settlement SLS table (if computed)
        5. Verdict paragraph

    :param filepath:  Output .docx file path.
    :param analysis:  Result dict from api.run_foundation_analysis().
    :param project:   Project name for the header.
    :param job_ref:   Job reference number.
    :param calc_by:   Initials of the analyst.
    :param checked_by: Initials of the checker.

    Reference: EC7 EN 1997-1:2004 §6, Annex D.
    """
    doc  = Document()
    _apply_normal_style(doc)

    section = doc.sections[0]
    section.page_width    = Cm(21.0)
    section.page_height   = Cm(29.7)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    date_str = datetime.date.today().strftime("%d %b %Y")

    # ── Header ───────────────────────────────────────────────────────────────
    _header_table(doc, project, job_ref, calc_by, checked_by, date_str)

    # ── 1. Input Parameters ──────────────────────────────────────────────────
    _section_heading(doc, "1.  Input Parameters  (EC7 §6)")
    s = analysis["soil"]
    f = analysis["foundation"]
    inp_rows = [
        ("Soil name",           s["name"],                     ""),
        ("Unit weight γ",       f"{s['gamma']:.1f}",           "kN/m³"),
        ("Friction angle φ'k",  f"{s['phi_k']:.1f}",           "°"),
        ("Cohesion c'k",        f"{s['c_k']:.1f}",             "kPa"),
        ("Width B",             f"{f['B']:.2f}",               "m"),
        ("Length L",            str(f["L"]) if f["L"] else "∞ (strip)", "m"),
        ("Embedment Df",        f"{f['Df']:.2f}",              "m"),
        ("Effective width B'",  f"{f['B_eff']:.3f}",           "m"),
        ("Effective area A'",   f"{f['A_eff']:.4f}",           "m²/m"),
    ]
    tbl = doc.add_table(rows=1 + len(inp_rows), cols=3)
    tbl.style = "Table Grid"
    tbl.autofit = False
    for i, w in enumerate([Cm(8), Cm(5), Cm(4)]):
        for cell in tbl.columns[i].cells:
            cell.width = w
    # Header row
    for txt, cell in zip(["Parameter", "Value", "Unit"], tbl.rows[0].cells):
        _set_cell_bg(cell, _C_HEADER_BG)
        _cell_para(cell, txt, bold=True, fg=_C_HEADER_FG)
        _set_cell_borders(cell)
    # Data rows
    for r_idx, (lbl, val, unit) in enumerate(inp_rows, start=1):
        bg = "EEEEEE" if r_idx % 2 == 0 else "FFFFFF"
        row = tbl.rows[r_idx]
        for txt, cell in zip([lbl, val, unit], row.cells):
            _set_cell_bg(cell, bg)
            _cell_para(cell, txt, bold=(txt == lbl))
            _set_cell_borders(cell)
    doc.add_paragraph()

    # ── 2. EC7 DA1 Bearing Capacity ──────────────────────────────────────────
    _section_heading(doc, "2.  EC7 DA1 Bearing Capacity — GEO ULS  (Annex D)")
    uls_hdrs = ["Combination", "γ_G", "γ_Q", "V_d (kN/m)", "R_d (kN/m)",
                "Utilisation", "Result"]
    tbl2 = doc.add_table(rows=1 + 2, cols=len(uls_hdrs))
    tbl2.style = "Table Grid"
    for txt, cell in zip(uls_hdrs, tbl2.rows[0].cells):
        _set_cell_bg(cell, _C_HEADER_BG)
        _cell_para(cell, txt, bold=True, fg=_C_HEADER_FG,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_borders(cell)
    for r_idx, c in enumerate([analysis["comb1"], analysis["comb2"]], start=1):
        ok_str = "PASS ✓" if c["passes"] else "FAIL ✗"
        fg_col = _C_PASS if c["passes"] else _C_FAIL
        bg = "EEEEEE" if r_idx % 2 == 0 else "FFFFFF"
        vals = [c["label"], f"{c['gG']:.2f}", f"{c['gQ']:.2f}",
                f"{c['Vd']:.1f}", f"{c['Rd']:.1f}",
                f"{c['utilisation']:.3f}", ok_str]
        for col_i, (txt, cell) in enumerate(zip(vals, tbl2.rows[r_idx].cells)):
            _set_cell_bg(cell, bg)
            is_result = (col_i == len(vals) - 1)
            _cell_para(cell, txt, bold=is_result,
                       fg=fg_col if is_result else None,
                       align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_borders(cell)
    doc.add_paragraph()

    # ── 3. Settlement SLS ────────────────────────────────────────────────────
    if analysis.get("s_total_mm") is not None:
        _section_heading(doc, "3.  Settlement Check — SLS  (EC7 §6.6)")
        sls_pass = analysis.get("sls_passes", False)
        sls_rows = [
            ("Immediate settlement s_i",     f"{analysis.get('s_immediate_mm', '—')} mm"),
            ("Consolidation settlement s_c",  f"{analysis.get('s_consolidation_mm', '—')} mm"),
            ("Total settlement s_total",      f"{analysis['s_total_mm']:.1f} mm"),
            ("SLS limit s_lim",               f"{analysis['s_lim_mm']:.1f} mm"),
            ("SLS result",                    "PASS ✓" if sls_pass else "FAIL ✗"),
        ]
        if analysis.get("t_95_years") is not None:
            sls_rows.append(("Time to 95% consolidation",
                             f"{analysis['t_95_years']:.1f} years"))
        tbl3 = doc.add_table(rows=len(sls_rows), cols=2)
        tbl3.style = "Table Grid"
        for r_idx, (lbl, val) in enumerate(sls_rows):
            bg   = "EEEEEE" if r_idx % 2 == 0 else "FFFFFF"
            is_r = (lbl == "SLS result")
            fg   = (_C_PASS if sls_pass else _C_FAIL) if is_r else None
            for txt, cell in zip([lbl, val], tbl3.rows[r_idx].cells):
                _set_cell_bg(cell, bg)
                _cell_para(cell, txt, bold=is_r, fg=fg)
                _set_cell_borders(cell)
        doc.add_paragraph()

    # ── 4. Verdict ───────────────────────────────────────────────────────────
    _section_heading(doc, "4.  Verdict")
    passes  = analysis.get("passes", False)
    verdict = "SATISFACTORY — PASS ✓" if passes else "UNSATISFACTORY — FAIL ✗"
    vp  = doc.add_paragraph()
    run = vp.add_run(verdict)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string(_C_PASS if passes else _C_FAIL)
    vp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Footer ───────────────────────────────────────────────────────────────
    doc.add_paragraph()
    fp = doc.add_paragraph(
        "Generated by DesignApp · Eurocode 7 EN 1997-1:2004 · "
        "Verify all inputs before use for design."
    )
    fp.runs[0].font.size = Pt(7)
    fp.runs[0].italic    = True
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if analysis.get("warnings"):
        _section_heading(doc, "Notes / Warnings")
        for w in analysis["warnings"]:
            doc.add_paragraph(f"• {w}", style="Normal")

    doc.save(filepath)


# ============================================================
#  Retaining wall DOCX report  (B-07)
# ============================================================

def generate_wall_report_docx(
    filepath    : str,
    analysis    : dict,
    project     : str = "Untitled Project",
    job_ref     : str = "—",
    calc_by     : str = "DesignApp",
    checked_by  : str = "—",
) -> None:
    """
    Editable Word (.docx) calculation sheet for EC7 DA1 retaining wall.

    Sections:
        1. Project header block
        2. Soil and wall geometry input table
        3. EC7 DA1 combinations — sliding, overturning, bearing
        4. Verdict paragraph

    :param filepath:  Output .docx file path.
    :param analysis:  Result dict from api.run_wall_analysis().
    :param project:   Project name for the header.
    :param job_ref:   Job reference number.
    :param calc_by:   Initials of the analyst.
    :param checked_by: Initials of the checker.

    Reference: EC7 EN 1997-1:2004 §9. Craig Ch.11. Bond & Harris Ch.14.
    """
    doc  = Document()
    _apply_normal_style(doc)

    section = doc.sections[0]
    section.page_width    = Cm(21.0)
    section.page_height   = Cm(29.7)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    date_str = datetime.date.today().strftime("%d %b %Y")

    # ── Header ───────────────────────────────────────────────────────────────
    _header_table(doc, project, job_ref, calc_by, checked_by, date_str)

    # ── 1. Input Parameters ──────────────────────────────────────────────────
    _section_heading(doc, "1.  Input Parameters  (EC7 §9)")
    s  = analysis["soil"]
    fs = analysis.get("foundation_soil", s)
    w  = analysis["wall"]
    inp_rows = [
        ("Backfill name",             s["name"],                    ""),
        ("Backfill γ / φ'k / c'k",   f"{s['gamma']:.1f} / {s['phi_k']:.1f} / {s['c_k']:.1f}", "kN/m³ / ° / kPa"),
        ("Foundation soil φ'k",       f"{fs['phi_k']:.1f}",         "°"),
        ("Retained height H",         f"{w['H_wall']:.2f}",         "m"),
        ("Base width B",              f"{w['B_base']:.2f}",         "m"),
        ("Toe projection b_toe",      f"{w['B_toe']:.2f}",          "m"),
        ("Heel projection b_heel",    f"{w.get('b_heel', '—')}",    "m"),
        ("Stem thickness (base/top)", f"{w['t_stem_base']:.2f} / {w['t_stem_top']:.2f}", "m"),
        ("Base slab thickness",       f"{w['t_base']:.2f}",         "m"),
        ("Ka (Rankine)",              f"{analysis['Ka']:.4f}",      "—"),
        ("Kp (Rankine)",              f"{analysis['Kp']:.4f}",      "—"),
    ]
    tbl = doc.add_table(rows=1 + len(inp_rows), cols=3)
    tbl.style = "Table Grid"
    for txt, cell in zip(["Parameter", "Value", "Unit"], tbl.rows[0].cells):
        _set_cell_bg(cell, _C_HEADER_BG)
        _cell_para(cell, txt, bold=True, fg=_C_HEADER_FG)
        _set_cell_borders(cell)
    for r_idx, (lbl, val, unit) in enumerate(inp_rows, start=1):
        bg = "EEEEEE" if r_idx % 2 == 0 else "FFFFFF"
        for txt, cell in zip([lbl, val, unit], tbl.rows[r_idx].cells):
            _set_cell_bg(cell, bg)
            _cell_para(cell, txt)
            _set_cell_borders(cell)
    doc.add_paragraph()

    # ── 2. EC7 DA1 Combinations ──────────────────────────────────────────────
    _section_heading(doc, "2.  EC7 DA1 ULS — Sliding, Overturning & Bearing  (§9)")
    da1_hdrs = ["Comb", "Ka", "Pa\n(kN/m)", "Slide\nFoS_d", "Slide",
                "Overt\nFoS_d", "Overt", "Bear η", "Bear", "Overall"]
    tbl2 = doc.add_table(rows=1 + 2, cols=len(da1_hdrs))
    tbl2.style = "Table Grid"
    for txt, cell in zip(da1_hdrs, tbl2.rows[0].cells):
        _set_cell_bg(cell, _C_HEADER_BG)
        _cell_para(cell, txt, bold=True, fg=_C_HEADER_FG,
                   align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=8.0)
        _set_cell_borders(cell)
    for r_idx, c in enumerate([analysis["comb1"], analysis["comb2"]], start=1):
        sl = c["sliding"]; ov = c["overturn"]; br = c["bearing"]
        ok_str = "PASS ✓" if c["passes"] else "FAIL ✗"
        fg     = _C_PASS if c["passes"] else _C_FAIL
        bg     = "EEEEEE" if r_idx % 2 == 0 else "FFFFFF"
        vals   = [
            c["label"], f"{c['ka']:.4f}", f"{c['Pa']:.1f}",
            f"{sl['fos_d']:.3f}", "✓" if sl["passes"] else "✗",
            f"{ov['fos_d']:.3f}", "✓" if ov["passes"] else "✗",
            f"{br['utilisation']:.3f}", "✓" if br["passes"] else "✗",
            ok_str,
        ]
        for col_i, (txt, cell) in enumerate(zip(vals, tbl2.rows[r_idx].cells)):
            _set_cell_bg(cell, bg)
            is_last = (col_i == len(vals) - 1)
            _cell_para(cell, txt, bold=is_last,
                       fg=fg if is_last else None,
                       align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=8.5)
            _set_cell_borders(cell)
    doc.add_paragraph()

    # ── 3. Verdict ───────────────────────────────────────────────────────────
    _section_heading(doc, "3.  Verdict")
    passes  = analysis.get("passes", False)
    verdict = "SATISFACTORY — PASS ✓" if passes else "UNSATISFACTORY — FAIL ✗"
    vp  = doc.add_paragraph()
    run = vp.add_run(verdict)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string(_C_PASS if passes else _C_FAIL)
    vp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Footer ───────────────────────────────────────────────────────────────
    doc.add_paragraph()
    fp = doc.add_paragraph(
        "Generated by DesignApp · Eurocode 7 EN 1997-1:2004 · "
        "Verify all inputs before use for design."
    )
    fp.runs[0].font.size = Pt(7)
    fp.runs[0].italic    = True
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if analysis.get("warnings"):
        _section_heading(doc, "Notes / Warnings")
        for w in analysis["warnings"]:
            doc.add_paragraph(f"• {w}", style="Normal")

    doc.save(filepath)


def generate_sheet_pile_report_docx(
    filepath    : str,
    analysis    : dict,
    project     : str = "Untitled Project",
    job_ref     : str = "—",
    calc_by     : str = "DesignApp",
    checked_by  : str = "—",
) -> None:
    """
    Editable Word (.docx) calculation sheet for EC7 DA1 sheet pile analysis.
    """
    doc = Document()
    _apply_normal_style(doc)

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    date_str = datetime.date.today().strftime("%d %b %Y")
    _header_table(doc, project, job_ref, calc_by, checked_by, date_str)

    wall = analysis.get("wall", {})

    _section_heading(doc, "1.  Input Parameters  (EC7 §9)")
    inp_rows = [
        ("Wall label", wall.get("label", "Sheet Pile"), ""),
        ("Retained height h", f"{wall.get('h_retained', '—')}", "m"),
        ("Support type", wall.get("support", "propped"), ""),
        ("Prop depth z_prop", f"{wall.get('z_prop', '—')}", "m"),
        ("Characteristic Ka / Kp", f"{analysis.get('Ka_k', '—')} / {analysis.get('Kp_k', '—')}", "—"),
    ]
    tbl = doc.add_table(rows=1 + len(inp_rows), cols=3)
    tbl.style = "Table Grid"
    for txt, cell in zip(["Parameter", "Value", "Unit"], tbl.rows[0].cells):
        _set_cell_bg(cell, _C_HEADER_BG)
        _cell_para(cell, txt, bold=True, fg=_C_HEADER_FG)
        _set_cell_borders(cell)
    for r_idx, (lbl, val, unit) in enumerate(inp_rows, start=1):
        bg = "EEEEEE" if r_idx % 2 == 0 else "FFFFFF"
        for txt, cell in zip([lbl, str(val), unit], tbl.rows[r_idx].cells):
            _set_cell_bg(cell, bg)
            _cell_para(cell, txt)
            _set_cell_borders(cell)
    doc.add_paragraph()

    _section_heading(doc, "2.  EC7 DA1 Results")
    da1_hdrs = ["Comb", "phi_d (deg)", "Ka_d", "Kp_d", "d_min (m)", "T (kN/m)", "M_max", "Overall"]
    tbl2 = doc.add_table(rows=3, cols=len(da1_hdrs))
    tbl2.style = "Table Grid"
    for txt, cell in zip(da1_hdrs, tbl2.rows[0].cells):
        _set_cell_bg(cell, _C_HEADER_BG)
        _cell_para(cell, txt, bold=True, fg=_C_HEADER_FG,
                   align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=8.0)
        _set_cell_borders(cell)

    for r_idx, comb in enumerate([analysis.get("comb1", {}), analysis.get("comb2", {})], start=1):
        bg = "EEEEEE" if r_idx % 2 == 0 else "FFFFFF"
        is_gov = comb.get("label") == analysis.get("governing")
        vals = [
            comb.get("label", "—"),
            f"{comb.get('phi_d_deg', '—')}",
            f"{comb.get('Ka_d', '—')}",
            f"{comb.get('Kp_d', '—')}",
            f"{comb.get('d_min', '—')}",
            f"{comb.get('T_k', '—')}",
            f"{comb.get('M_max', '—')}",
            "GOV" if is_gov else "",
        ]
        for col_i, (txt, cell) in enumerate(zip(vals, tbl2.rows[r_idx].cells)):
            _set_cell_bg(cell, bg)
            _cell_para(cell, txt,
                       bold=(col_i == len(vals) - 1 and bool(txt)),
                       fg=_C_PASS if col_i == len(vals) - 1 and bool(txt) else None,
                       align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=8.5)
            _set_cell_borders(cell)
    doc.add_paragraph()

    _section_heading(doc, "3.  Verdict")
    passes = analysis.get("passes", False)
    verdict = "SATISFACTORY - PASS" if passes else "UNSATISFACTORY - FAIL"
    vp = doc.add_paragraph()
    run = vp.add_run(verdict)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string(_C_PASS if passes else _C_FAIL)
    vp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if analysis.get("warnings"):
        _section_heading(doc, "Notes / Warnings")
        for warning in analysis["warnings"]:
            doc.add_paragraph(f"• {warning}", style="Normal")

    doc.save(filepath)
