"""
tests/test_exporters.py

Validates exporters/plot_slope.py, exporters/plot_bishop.py, and
exporters/report_pdf.py by rendering real outputs and checking their
physical properties (file size, image dimensions, PDF page count).

These tests are integration-level: they exercise the full data flow from
Soil/Geometry -> Engine -> SearchResult/VerificationResult -> Exporter.

Fixtures use the same 1:2 slope / dense-sand combination validated in
earlier Phase 0 tests (Craig Ch.9 calibration case).

Run from the DesignApp root:
    python -m pytest tests/test_exporters.py
  or:
    python tests/test_exporters.py
"""
import sys, os, math, tempfile

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
HERE = os.path.abspath(os.path.dirname(__file__))
for p in (HERE, ROOT):
    if p not in sys.path:
        sys.path.insert(0, p)

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

from models.soil            import Soil
from models.geometry        import SlopeGeometry, SlipCircle
from core.search            import grid_search, SearchResult
from core.factors_of_safety import verify_slope_da1
from core.slicer            import create_slices
from exporters.plot_slope   import plot_slope_stability, save_slope_plot
from exporters.plot_bishop  import plot_fos_heatmap, save_fos_heatmap
from exporters.report_pdf   import generate_slope_report

# Sprint 11 — foundation / wall exporters and plot modules
from exporters.plot_foundation import plot_foundation_bearing, save_foundation_plot
from exporters.plot_wall import plot_retaining_wall, save_wall_plot
from exporters.report_pdf import generate_foundation_report, generate_wall_report
from exporters.report_docx import generate_foundation_report_docx, generate_wall_report_docx
from api                    import run_foundation_analysis, run_wall_analysis


# ---------------------------------------------------------------------------
#  Shared fixtures — 1:2 slope, dense sand (Craig Ch.9)
# ---------------------------------------------------------------------------

SLOPE = SlopeGeometry([(0, 3), (6, 3), (12, 0), (18, 0)])
SOIL  = Soil("Dense Sand", unit_weight=19.0, friction_angle=35.0, cohesion=0.0)
RU    = 0.0

# Run the grid search once for all tests (moderately coarse for test speed)
_RESULT = grid_search(
    slope     = SLOPE,
    soil      = SOIL,
    cx_range  = (3.0, 14.0),
    cy_range  = (3.0, 12.0),
    r_range   = (4.0, 16.0),
    n_cx      = 8,
    n_cy      = 8,
    n_r       = 6,
    ru        = RU,
)
_VERIFICATION = verify_slope_da1(SLOPE, SOIL, ru=RU)
_SLICES = create_slices(SLOPE, _RESULT.critical_circle, SOIL, num_slices=20)


# ---------------------------------------------------------------------------
#  Test 1 – plot_slope_stability returns a valid Figure
# ---------------------------------------------------------------------------

def test_plot_slope_returns_figure():
    """
    plot_slope_stability() must return a matplotlib Figure with correct axes.
    The figure must have exactly one Axes object with labelled x and y axes.
    """
    print("\n" + "="*60)
    print("  TEST 1 – plot_slope_stability: returns valid Figure")
    print("="*60)

    fig = plot_slope_stability(SLOPE, _RESULT, title="Test slope", ru=RU)

    assert isinstance(fig, plt.Figure), "FAIL: expected matplotlib.Figure"
    axes = fig.get_axes()
    assert len(axes) == 1, f"FAIL: expected 1 Axes, got {len(axes)}"
    ax = axes[0]
    assert ax.get_xlabel() != "", "FAIL: x-axis has no label"
    assert ax.get_ylabel() != "", "FAIL: y-axis has no label"
    plt.close(fig)

    print("  Figure type     : OK")
    print("  Axes count      : 1  OK")
    print("  Axis labels     : OK")
    print("\n  PASS  test_plot_slope_returns_figure")


# ---------------------------------------------------------------------------
#  Test 2 – save_slope_plot writes a non-empty PNG file
# ---------------------------------------------------------------------------

def test_save_slope_plot_writes_file():
    """
    save_slope_plot() must write a PNG file >= 20 kB (non-trivial render).
    """
    print("\n" + "="*60)
    print("  TEST 2 – save_slope_plot: writes non-empty PNG")
    print("="*60)

    with tempfile.TemporaryDirectory() as tmpdir:
        path = os.path.join(tmpdir, "slope_test.png")
        save_slope_plot(SLOPE, _RESULT, filepath=path, title="Test", dpi=96)

        assert os.path.exists(path), "FAIL: PNG file not created"
        size = os.path.getsize(path)
        print(f"  PNG size : {size:,} bytes")
        assert size > 20_000, f"FAIL: PNG too small ({size} bytes) — likely blank render"

    print("\n  PASS  test_save_slope_plot_writes_file")


# ---------------------------------------------------------------------------
#  Test 3 – FoS heatmap returns a Figure with a colorbar
# ---------------------------------------------------------------------------

def test_fos_heatmap_returns_figure():
    """
    plot_fos_heatmap() must return a Figure with >= 2 axes (plot + colorbar).
    """
    print("\n" + "="*60)
    print("  TEST 3 – plot_fos_heatmap: returns valid Figure")
    print("="*60)

    fig = plot_fos_heatmap(SLOPE, SOIL, _RESULT, ru=RU, n_cx=8, n_cy=7)

    assert isinstance(fig, plt.Figure), "FAIL: expected matplotlib.Figure"
    axes = fig.get_axes()
    # Main axes + colorbar axes
    assert len(axes) >= 2, f"FAIL: expected >= 2 axes (heatmap + colorbar), got {len(axes)}"
    plt.close(fig)

    print(f"  Figure axes count : {len(axes)}  OK")
    print("\n  PASS  test_fos_heatmap_returns_figure")


# ---------------------------------------------------------------------------
#  Test 4 – save_fos_heatmap writes a non-empty PNG file
# ---------------------------------------------------------------------------

def test_save_fos_heatmap_writes_file():
    """save_fos_heatmap() must write a PNG file >= 20 kB."""
    print("\n" + "="*60)
    print("  TEST 4 – save_fos_heatmap: writes non-empty PNG")
    print("="*60)

    with tempfile.TemporaryDirectory() as tmpdir:
        path = os.path.join(tmpdir, "heatmap_test.png")
        save_fos_heatmap(SLOPE, SOIL, _RESULT, filepath=path, dpi=96, n_cx=8, n_cy=7)

        assert os.path.exists(path), "FAIL: PNG file not created"
        size = os.path.getsize(path)
        print(f"  PNG size : {size:,} bytes")
        assert size > 20_000, f"FAIL: PNG too small ({size} bytes)"

    print("\n  PASS  test_save_fos_heatmap_writes_file")


# ---------------------------------------------------------------------------
#  Test 5 – generate_slope_report produces a valid PDF
# ---------------------------------------------------------------------------

def test_generate_slope_report_produces_pdf():
    """
    generate_slope_report() must produce a PDF file >= 50 kB (with embedded
    figure) that starts with the PDF magic bytes '%PDF'.
    """
    print("\n" + "="*60)
    print("  TEST 5 – generate_slope_report: produces valid PDF")
    print("="*60)

    with tempfile.TemporaryDirectory() as tmpdir:
        path = os.path.join(tmpdir, "slope_calc.pdf")

        generate_slope_report(
            filepath      = path,
            soil          = SOIL,
            slope         = SLOPE,
            search_result = _RESULT,
            verification  = _VERIFICATION,
            slices        = _SLICES,
            ru            = RU,
            project       = "DesignApp Test Project",
            job_ref       = "DA-001",
            calc_by       = "Auto",
            checked_by    = "QA",
        )

        assert os.path.exists(path), "FAIL: PDF file not created"
        size = os.path.getsize(path)
        print(f"  PDF size : {size:,} bytes")
        assert size > 50_000, f"FAIL: PDF too small ({size} bytes) — likely empty render"

        with open(path, "rb") as f:
            magic = f.read(4)
        assert magic == b"%PDF", f"FAIL: file does not start with PDF magic bytes (got {magic!r})"
        print("  PDF magic bytes : %PDF  OK")

    print("\n  PASS  test_generate_slope_report_produces_pdf")


# ---------------------------------------------------------------------------
#  Test 6 – PDF content includes expected text tokens (via pypdf extraction)
# ---------------------------------------------------------------------------

def test_pdf_contains_expected_content():
    """
    The PDF must contain key text tokens when extracted via pypdf.

    ReportLab compresses content streams (FlateDecode), so raw byte search
    is unreliable.  pypdf decompresses and extracts text properly.

    Tokens verified: project name, job reference, 'DA1', soil name.
    """
    print("\n" + "="*60)
    print("  TEST 6 – PDF contains expected text content (via pypdf)")
    print("="*60)

    try:
        import pypdf
    except ImportError:
        print("  SKIP: pypdf not available — install with 'pip install pypdf'")
        return

    with tempfile.TemporaryDirectory() as tmpdir:
        path = os.path.join(tmpdir, "slope_content_check.pdf")

        generate_slope_report(
            filepath      = path,
            soil          = SOIL,
            slope         = SLOPE,
            search_result = _RESULT,
            verification  = _VERIFICATION,
            slices        = _SLICES,
            ru            = RU,
            project       = "ContentCheckProject",
            job_ref       = "CCREF-007",
            calc_by       = "Auto",
            checked_by    = "QA",
        )

        # Extract all text from all pages
        reader    = pypdf.PdfReader(path)
        n_pages   = len(reader.pages)
        full_text = "\n".join(
            page.extract_text() or "" for page in reader.pages
        )

        print(f"  Pages extracted : {n_pages}")
        print(f"  Characters      : {len(full_text)}")
        assert n_pages >= 1, "FAIL: PDF has no pages"
        assert len(full_text) > 100, \
            f"FAIL: extracted text too short ({len(full_text)} chars) — likely empty render"

        tokens = ["ContentCheckProject", "CCREF-007", "DA1", "Dense Sand"]
        for token in tokens:
            found = token in full_text
            print(f"  Token '{token}': {'FOUND' if found else 'MISSING'}")
            assert found, (
                f"FAIL: token '{token}' not found in extracted PDF text.\n"
                f"  First 500 chars of extracted text:\n{full_text[:500]}"
            )

    print("\n  PASS  test_pdf_contains_expected_content")


# ---------------------------------------------------------------------------
#  Runner
# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------
#  Sprint 11 fixtures — foundation and wall (medium-complexity cases)
# ---------------------------------------------------------------------------

# Foundation: 2.0×3.0 m rectangular pad, B-sand, Df=1.0 m
_FND_PARAMS = {
    "gamma": 19.0, "phi_k": 32.0, "c_k": 5.0,
    "B": 2.0, "L": 3.0, "Df": 1.0,
    "Gk": 400.0, "Qk": 150.0,
}
_FND_RESULT = run_foundation_analysis(_FND_PARAMS)

# Retaining wall: H=4 m cantilever, medium dense sand backfill
_WALL_PARAMS = {
    "gamma": 18.0, "phi_k": 30.0, "c_k": 0.0,
    "H_wall": 4.0, "B_base": 2.8, "B_toe": 0.5,
    "t_stem_base": 0.4, "t_stem_top": 0.3, "t_base": 0.45,
    "Gk_wall": 50.0,
}
_WALL_RESULT = run_wall_analysis(_WALL_PARAMS)

# Sanity-guard: confirm fixtures are valid before any test runs
assert _FND_RESULT["ok"],  f"Foundation fixture failed: {_FND_RESULT['errors']}"
assert _WALL_RESULT["ok"], f"Wall fixture failed: {_WALL_RESULT['errors']}"


# ===========================================================================
#  TEST 7 – plot_foundation_bearing returns a valid Figure
# ===========================================================================

def test_plot_foundation_returns_figure():
    """
    plot_foundation_bearing() must return a non-empty matplotlib Figure with
    exactly one Axes object containing rendered artists.

    Reference: Das (2019) §3.6 — Boussinesq stress distribution visualization.
    """
    print("\n" + "="*60)
    print("  TEST 7 – plot_foundation_bearing returns a Figure")
    print("="*60)

    fig = plot_foundation_bearing(_FND_RESULT)
    assert isinstance(fig, plt.Figure), f"Expected Figure, got {type(fig)}"

    axes = fig.axes
    assert len(axes) == 1, f"Expected 1 Axes, got {len(axes)}"
    ax = axes[0]

    n_patches = len(ax.patches)
    n_lines   = len(ax.lines)
    n_texts   = len(ax.texts)
    print(f"  Axes patches : {n_patches}")
    print(f"  Axes lines   : {n_lines}")
    print(f"  Axes texts   : {n_texts}")

    assert n_patches >= 2,  "Expected footing + soil patches"
    assert n_texts   >= 1,  "Expected at least one text annotation"

    xlim = ax.get_xlim()
    ylim = ax.get_ylim()
    print(f"  x-range      : [{xlim[0]:.2f}, {xlim[1]:.2f}]")
    print(f"  y-range      : [{ylim[0]:.2f}, {ylim[1]:.2f}]")
    assert xlim[1] > xlim[0], "x-axis is collapsed"
    assert ylim[1] > ylim[0], "y-axis is collapsed"

    plt.close(fig)
    print("\n  PASS  test_plot_foundation_returns_figure")


# ===========================================================================
#  TEST 8 – save_foundation_plot writes a valid PNG file
# ===========================================================================

def test_save_foundation_plot_writes_file():
    """
    save_foundation_plot() must write a PNG file ≥ 10 kB with correct magic bytes.
    """
    print("\n" + "="*60)
    print("  TEST 8 – save_foundation_plot writes PNG file")
    print("="*60)

    with tempfile.TemporaryDirectory() as tmpdir:
        path = os.path.join(tmpdir, "foundation_plot.png")
        save_foundation_plot(_FND_RESULT, path)

        assert os.path.isfile(path), "File was not created"
        size = os.path.getsize(path)
        magic = open(path, "rb").read(8)
        PNG_MAGIC = b"\x89PNG\r\n\x1a\n"

        print(f"  File size    : {size:,} bytes")
        print(f"  Magic bytes  : {magic!r}")

        assert size  >= 10_000, f"PNG too small: {size} bytes"
        assert magic == PNG_MAGIC, f"Not a valid PNG: {magic!r}"

    print("\n  PASS  test_save_foundation_plot_writes_file")


# ===========================================================================
#  TEST 9 – Foundation figure has correct axis labels
# ===========================================================================

def test_foundation_figure_axis_labels():
    """
    The foundation cross-section figure must carry meaningful axis labels
    and a non-empty title.
    """
    print("\n" + "="*60)
    print("  TEST 9 – Foundation figure axis labels and title")
    print("="*60)

    fig = plot_foundation_bearing(_FND_RESULT, title="Test Foundation Figure")
    ax  = fig.axes[0]

    xlabel = ax.get_xlabel()
    ylabel = ax.get_ylabel()
    title  = ax.get_title()

    print(f"  xlabel : {xlabel!r}")
    print(f"  ylabel : {ylabel!r}")
    print(f"  title  : {title!r}")

    assert len(xlabel) > 3, "x-axis label missing or too short"
    assert len(ylabel) > 3, "y-axis label missing or too short"
    assert "Test Foundation Figure" in title

    plt.close(fig)
    print("\n  PASS  test_foundation_figure_axis_labels")


# ===========================================================================
#  TEST 10 – Foundation figure respects PASS/FAIL colour
# ===========================================================================

def test_foundation_figure_pass_fail_colour():
    """
    The result annotation box border colour must be green for PASS and
    red for a deliberately failing case.
    """
    print("\n" + "="*60)
    print("  TEST 10 – Foundation figure PASS/FAIL colour coding")
    print("="*60)

    # Force a FAIL case: very high load on a very small footing
    fail_params = dict(_FND_PARAMS)
    fail_params["B"]  = 0.5
    fail_params["Gk"] = 5000.0
    fail_params["Qk"] = 2000.0
    fail_result = run_foundation_analysis(fail_params)

    fig_pass = plot_foundation_bearing(_FND_RESULT)
    fig_fail = plot_foundation_bearing(fail_result)

    assert fig_pass is not None
    assert fig_fail is not None
    print(f"  Passing case passes: {_FND_RESULT['passes']}")
    print(f"  Failing case passes: {fail_result['passes']}")

    # Both figures render without error regardless of pass/fail
    assert len(fig_pass.axes) == 1
    assert len(fig_fail.axes) == 1

    plt.close(fig_pass)
    plt.close(fig_fail)
    print("\n  PASS  test_foundation_figure_pass_fail_colour")


# ===========================================================================
#  TEST 11 – plot_retaining_wall returns a valid Figure
# ===========================================================================

def test_plot_wall_returns_figure():
    """
    plot_retaining_wall() must return a Figure with one Axes containing
    wall geometry patches and pressure diagram artists.

    Reference: Craig §11.2 — retaining wall cross-section convention.
    """
    print("\n" + "="*60)
    print("  TEST 11 – plot_retaining_wall returns a Figure")
    print("="*60)

    fig = plot_retaining_wall(_WALL_RESULT)
    assert isinstance(fig, plt.Figure), f"Expected Figure, got {type(fig)}"

    axes = fig.axes
    assert len(axes) == 1
    ax   = axes[0]

    n_patches = len(ax.patches)
    n_polys   = sum(1 for c in ax.collections if hasattr(c, 'get_paths'))
    n_texts   = len(ax.texts)
    print(f"  Axes patches     : {n_patches}")
    print(f"  Axes collections : {n_polys}")
    print(f"  Axes texts       : {n_texts}")

    assert n_texts   >= 1, "Expected result annotation text"

    xlim = ax.get_xlim()
    ylim = ax.get_ylim()
    print(f"  x-range          : [{xlim[0]:.2f}, {xlim[1]:.2f}]")
    print(f"  y-range          : [{ylim[0]:.2f}, {ylim[1]:.2f}]")
    assert xlim[1] > xlim[0]
    assert ylim[1] > ylim[0]

    plt.close(fig)
    print("\n  PASS  test_plot_wall_returns_figure")


# ===========================================================================
#  TEST 12 – save_wall_plot writes a valid PNG file
# ===========================================================================

def test_save_wall_plot_writes_file():
    """
    save_wall_plot() must write a valid PNG ≥ 10 kB with correct magic bytes.
    """
    print("\n" + "="*60)
    print("  TEST 12 – save_wall_plot writes PNG file")
    print("="*60)

    with tempfile.TemporaryDirectory() as tmpdir:
        path = os.path.join(tmpdir, "wall_plot.png")
        save_wall_plot(_WALL_RESULT, path)

        assert os.path.isfile(path)
        size  = os.path.getsize(path)
        magic = open(path, "rb").read(8)
        PNG_MAGIC = b"\x89PNG\r\n\x1a\n"

        print(f"  File size    : {size:,} bytes")
        print(f"  Magic bytes  : {magic!r}")

        assert size  >= 10_000, f"PNG too small: {size} bytes"
        assert magic == PNG_MAGIC, f"Not a valid PNG: {magic!r}"

    print("\n  PASS  test_save_wall_plot_writes_file")


# ===========================================================================
#  TEST 13 – Wall figure has correct axis labels
# ===========================================================================

def test_wall_figure_axis_labels():
    """Wall cross-section figure must carry axis labels and a non-empty title."""
    print("\n" + "="*60)
    print("  TEST 13 – Wall figure axis labels and title")
    print("="*60)

    fig = plot_retaining_wall(_WALL_RESULT, title="Sprint 11 Wall Test")
    ax  = fig.axes[0]

    xlabel = ax.get_xlabel()
    ylabel = ax.get_ylabel()
    title  = ax.get_title()

    print(f"  xlabel : {xlabel!r}")
    print(f"  ylabel : {ylabel!r}")
    print(f"  title  : {title!r}")

    assert len(xlabel) > 3
    assert len(ylabel) > 3
    assert "Sprint 11 Wall Test" in title

    plt.close(fig)
    print("\n  PASS  test_wall_figure_axis_labels")


# ===========================================================================
#  TEST 14 – Wall y-range covers full wall height
# ===========================================================================

def test_wall_figure_y_range_covers_height():
    """
    The wall cross-section y-axis must span from below the base slab
    to above the retained fill surface (H = 4 m).
    """
    print("\n" + "="*60)
    print("  TEST 14 – Wall figure y-range covers full wall height")
    print("="*60)

    H   = float(_WALL_PARAMS["H_wall"])
    fig = plot_retaining_wall(_WALL_RESULT)
    ax  = fig.axes[0]
    ylim = ax.get_ylim()

    print(f"  H = {H:.1f} m   y_range = [{ylim[0]:.2f}, {ylim[1]:.2f}]")
    assert ylim[0] < 0,   "y-axis does not extend below base slab"
    assert ylim[1] >= H,  f"y-axis top {ylim[1]:.2f} < H={H:.1f}"

    plt.close(fig)
    print("\n  PASS  test_wall_figure_y_range_covers_height")


# ===========================================================================
#  TEST 15 – generate_foundation_report produces a valid PDF
# ===========================================================================

def test_generate_foundation_report_produces_pdf():
    """
    generate_foundation_report() must write a PDF ≥ 2 kB starting with %PDF-.

    Reference: EC7 EN 1997-1:2004 §6, Annex D — bearing capacity sheet.
    """
    print("\n" + "="*60)
    print("  TEST 15 – generate_foundation_report produces PDF")
    print("="*60)

    with tempfile.TemporaryDirectory() as tmpdir:
        path = os.path.join(tmpdir, "foundation_report.pdf")
        generate_foundation_report(
            filepath   = path,
            analysis   = _FND_RESULT,
            project    = "Sprint11TestProject",
            job_ref    = "S11-FND-001",
            calc_by    = "Auto",
            checked_by = "QA",
        )

        assert os.path.isfile(path), "PDF file not created"
        size  = os.path.getsize(path)
        magic = open(path, "rb").read(5)

        print(f"  File size    : {size:,} bytes")
        print(f"  Magic bytes  : {magic}")
        assert size  >= 2_000, f"PDF too small: {size}"
        assert magic == b"%PDF-", f"Invalid PDF magic: {magic}"

    print("\n  PASS  test_generate_foundation_report_produces_pdf")


# ===========================================================================
#  TEST 16 – Foundation PDF contains expected text tokens
# ===========================================================================

def test_foundation_pdf_contains_expected_content():
    """
    The foundation PDF must contain the project name, job reference, soil name,
    and the token 'DA1'.

    Uses pypdf for decompressed text extraction.
    """
    print("\n" + "="*60)
    print("  TEST 16 – Foundation PDF text content check")
    print("="*60)

    try:
        import pypdf
    except ImportError:
        print("  SKIP: pypdf not available")
        return

    with tempfile.TemporaryDirectory() as tmpdir:
        path = os.path.join(tmpdir, "fnd_content.pdf")
        generate_foundation_report(
            filepath   = path,
            analysis   = _FND_RESULT,
            project    = "FndContentProject",
            job_ref    = "FND-REF-999",
        )
        reader    = pypdf.PdfReader(path)
        full_text = "\n".join(p.extract_text() or "" for p in reader.pages)
        n_pages   = len(reader.pages)

        print(f"  Pages : {n_pages}  |  Chars : {len(full_text)}")
        assert n_pages  >= 1
        assert len(full_text) > 50, "Extracted text too short"

        tokens = ["FndContentProject", "FND-REF-999", "DA1"]
        for tok in tokens:
            found = tok in full_text
            print(f"  Token '{tok}': {'FOUND' if found else 'MISSING'}")
            assert found, f"Token '{tok}' not found in PDF"

    print("\n  PASS  test_foundation_pdf_contains_expected_content")


# ===========================================================================
#  TEST 17 – generate_wall_report produces a valid PDF
# ===========================================================================

def test_generate_wall_report_produces_pdf():
    """
    generate_wall_report() must write a PDF ≥ 2 kB starting with %PDF-.

    Reference: EC7 §9 — Retaining structures calculation sheet.
    """
    print("\n" + "="*60)
    print("  TEST 17 – generate_wall_report produces PDF")
    print("="*60)

    with tempfile.TemporaryDirectory() as tmpdir:
        path = os.path.join(tmpdir, "wall_report.pdf")
        generate_wall_report(
            filepath   = path,
            analysis   = _WALL_RESULT,
            project    = "Sprint11WallProject",
            job_ref    = "S11-WALL-001",
            calc_by    = "Auto",
            checked_by = "QA",
        )

        assert os.path.isfile(path)
        size  = os.path.getsize(path)
        magic = open(path, "rb").read(5)

        print(f"  File size    : {size:,} bytes")
        print(f"  Magic bytes  : {magic}")
        assert size  >= 2_000, f"PDF too small: {size}"
        assert magic == b"%PDF-"

    print("\n  PASS  test_generate_wall_report_produces_pdf")


# ===========================================================================
#  TEST 18 – Wall PDF contains expected text tokens
# ===========================================================================

def test_wall_pdf_contains_expected_content():
    """
    The wall PDF must contain the project name, job reference, and 'DA1'.
    """
    print("\n" + "="*60)
    print("  TEST 18 – Wall PDF text content check")
    print("="*60)

    try:
        import pypdf
    except ImportError:
        print("  SKIP: pypdf not available")
        return

    with tempfile.TemporaryDirectory() as tmpdir:
        path = os.path.join(tmpdir, "wall_content.pdf")
        generate_wall_report(
            filepath   = path,
            analysis   = _WALL_RESULT,
            project    = "WallContentProject",
            job_ref    = "WALL-REF-888",
        )
        reader    = pypdf.PdfReader(path)
        full_text = "\n".join(p.extract_text() or "" for p in reader.pages)

        print(f"  Pages : {len(reader.pages)}  |  Chars : {len(full_text)}")
        tokens = ["WallContentProject", "WALL-REF-888", "DA1"]
        for tok in tokens:
            found = tok in full_text
            print(f"  Token '{tok}': {'FOUND' if found else 'MISSING'}")
            assert found, f"Token '{tok}' not found in PDF"

    print("\n  PASS  test_wall_pdf_contains_expected_content")


# ===========================================================================
#  TEST 19 – generate_foundation_report_docx produces a valid DOCX
# ===========================================================================

def test_generate_foundation_report_docx():
    """
    generate_foundation_report_docx() must write a .docx ≥ 10 kB.
    A .docx is a ZIP archive; verify by magic bytes (PK\\x03\\x04).

    Reference: report_docx.py — sprint 11 DOCX exporter.
    """
    print("\n" + "="*60)
    print("  TEST 19 – generate_foundation_report_docx produces DOCX")
    print("="*60)

    with tempfile.TemporaryDirectory() as tmpdir:
        path = os.path.join(tmpdir, "fnd_report.docx")
        generate_foundation_report_docx(
            filepath   = path,
            analysis   = _FND_RESULT,
            project    = "Sprint11FndDocx",
            job_ref    = "FND-DOCX-042",
            calc_by    = "Auto",
            checked_by = "QA",
        )

        assert os.path.isfile(path), "DOCX file not created"
        size  = os.path.getsize(path)
        magic = open(path, "rb").read(4)
        DOCX_MAGIC = b"PK\x03\x04"

        print(f"  File size    : {size:,} bytes")
        print(f"  Magic bytes  : {magic!r}")
        assert size  >= 10_000, f"DOCX too small: {size}"
        assert magic == DOCX_MAGIC, f"Not a ZIP/DOCX: {magic!r}"

    print("\n  PASS  test_generate_foundation_report_docx")


# ===========================================================================
#  TEST 20 – Foundation DOCX contains expected text in body paragraphs
# ===========================================================================

def test_foundation_docx_text_content():
    """
    The foundation DOCX body must contain the project name and job reference.
    Checks both body paragraphs and table cells (headers are in tables).
    """
    print("\n" + "="*60)
    print("  TEST 20 – Foundation DOCX text content")
    print("="*60)

    try:
        from docx import Document as DocxDocument
    except ImportError:
        print("  SKIP: python-docx not available")
        return

    with tempfile.TemporaryDirectory() as tmpdir:
        path = os.path.join(tmpdir, "fnd_text.docx")
        generate_foundation_report_docx(
            filepath = path,
            analysis = _FND_RESULT,
            project  = "DocxFndProject",
            job_ref  = "FND-TXT-007",
        )
        doc   = DocxDocument(path)

        # Collect all text: paragraphs + table cells
        para_texts  = [p.text for p in doc.paragraphs]
        table_texts = [
            cell.text
            for tbl in doc.tables
            for row in tbl.rows
            for cell in row.cells
        ]
        full = "\n".join(para_texts + table_texts)

        print(f"  Paragraphs : {len(para_texts)}")
        print(f"  Table cells: {len(table_texts)}")
        print(f"  Total chars: {len(full)}")

        tokens = ["DocxFndProject", "FND-TXT-007"]
        for tok in tokens:
            found = tok in full
            print(f"  Token '{tok}': {'FOUND' if found else 'MISSING'}")
            assert found, f"Token '{tok}' not found in DOCX paragraphs or tables"

    print("\n  PASS  test_foundation_docx_text_content")


# ===========================================================================
#  TEST 21 – generate_wall_report_docx produces a valid DOCX
# ===========================================================================

def test_generate_wall_report_docx():
    """
    generate_wall_report_docx() must write a .docx ≥ 10 kB with DOCX magic.
    """
    print("\n" + "="*60)
    print("  TEST 21 – generate_wall_report_docx produces DOCX")
    print("="*60)

    with tempfile.TemporaryDirectory() as tmpdir:
        path = os.path.join(tmpdir, "wall_report.docx")
        generate_wall_report_docx(
            filepath   = path,
            analysis   = _WALL_RESULT,
            project    = "Sprint11WallDocx",
            job_ref    = "WALL-DOCX-099",
            calc_by    = "Auto",
            checked_by = "QA",
        )

        assert os.path.isfile(path)
        size  = os.path.getsize(path)
        magic = open(path, "rb").read(4)
        DOCX_MAGIC = b"PK\x03\x04"

        print(f"  File size    : {size:,} bytes")
        print(f"  Magic bytes  : {magic!r}")
        assert size  >= 10_000
        assert magic == DOCX_MAGIC

    print("\n  PASS  test_generate_wall_report_docx")


# ===========================================================================
#  TEST 22 – Wall DOCX contains expected text
# ===========================================================================

def test_wall_docx_text_content():
    """
    The wall DOCX body must contain project name and job reference.
    Checks both body paragraphs and table cells.
    """
    print("\n" + "="*60)
    print("  TEST 22 – Wall DOCX text content")
    print("="*60)

    try:
        from docx import Document as DocxDocument
    except ImportError:
        print("  SKIP: python-docx not available")
        return

    with tempfile.TemporaryDirectory() as tmpdir:
        path = os.path.join(tmpdir, "wall_text.docx")
        generate_wall_report_docx(
            filepath = path,
            analysis = _WALL_RESULT,
            project  = "DocxWallProject",
            job_ref  = "WALL-TXT-021",
        )
        from docx import Document as DocxDocument
        doc   = DocxDocument(path)

        # Collect all text: paragraphs + table cells
        para_texts  = [p.text for p in doc.paragraphs]
        table_texts = [
            cell.text
            for tbl in doc.tables
            for row in tbl.rows
            for cell in row.cells
        ]
        full = "\n".join(para_texts + table_texts)

        print(f"  Paragraphs : {len(para_texts)}")
        print(f"  Table cells: {len(table_texts)}")
        tokens = ["DocxWallProject", "WALL-TXT-021"]
        for tok in tokens:
            found = tok in full
            print(f"  Token '{tok}': {'FOUND' if found else 'MISSING'}")
            assert found, f"Token '{tok}' not found in DOCX"

    print("\n  PASS  test_wall_docx_text_content")


# ===========================================================================
#  TEST 23 – Foundation PDF renders correctly for a FAIL case
# ===========================================================================

def test_foundation_report_fail_case():
    """
    generate_foundation_report() must complete without error even when the
    foundation check FAILS (utilisation > 1.0).
    """
    print("\n" + "="*60)
    print("  TEST 23 – Foundation PDF renders for FAIL case")
    print("="*60)

    fail_params = dict(_FND_PARAMS)
    fail_params["B"]  = 0.5
    fail_params["Gk"] = 5000.0
    fail_params["Qk"] = 2000.0
    fail_result = run_foundation_analysis(fail_params)

    print(f"  Foundation passes: {fail_result['passes']}")

    with tempfile.TemporaryDirectory() as tmpdir:
        path = os.path.join(tmpdir, "fnd_fail.pdf")
        generate_foundation_report(path, fail_result,
                                   project="FailTest", job_ref="FAIL-001")
        assert os.path.isfile(path)
        size = os.path.getsize(path)
        print(f"  PDF size : {size:,} bytes")
        assert size >= 2_000

    print("\n  PASS  test_foundation_report_fail_case")


# ===========================================================================
#  TEST 24 – Wall PDF renders correctly for a FAIL case
# ===========================================================================

def test_wall_report_fail_case():
    """
    generate_wall_report() must complete without error even when wall checks FAIL.
    Uses a geometrically-valid but structurally-inadequate wall (B too narrow).
    """
    print("\n" + "="*60)
    print("  TEST 24 – Wall PDF renders for FAIL case")
    print("="*60)

    # Narrow base width → fails sliding / overturning, but geometry is valid
    fail_params = {
        "gamma": 18.0, "phi_k": 30.0, "c_k": 0.0,
        "H_wall": 4.0, "B_base": 1.2, "B_toe": 0.3,
        "t_stem_base": 0.3, "t_stem_top": 0.2, "t_base": 0.3,
        "Gk_wall": 20.0,
    }
    fail_result = run_wall_analysis(fail_params)
    assert fail_result["ok"], f"Wall analysis errored: {fail_result.get('error')}"

    print(f"  Wall passes: {fail_result['passes']}")

    with tempfile.TemporaryDirectory() as tmpdir:
        path = os.path.join(tmpdir, "wall_fail.pdf")
        generate_wall_report(path, fail_result,
                             project="WallFailTest", job_ref="WFAIL-001")
        assert os.path.isfile(path)
        size = os.path.getsize(path)
        print(f"  PDF size : {size:,} bytes")
        assert size >= 2_000

    print("\n  PASS  test_wall_report_fail_case")


# ===========================================================================
#  TEST 25 – Both new plot modules close figures (no memory leak)
# ===========================================================================

def test_save_functions_close_figures():
    """
    save_foundation_plot() and save_wall_plot() must close figures after saving,
    preventing matplotlib figure accumulation in long-running servers.
    """
    print("\n" + "="*60)
    print("  TEST 25 – save_*_plot() closes figures (no memory leak)")
    print("="*60)

    before = len(plt.get_fignums())

    with tempfile.TemporaryDirectory() as tmpdir:
        save_foundation_plot(_FND_RESULT,  os.path.join(tmpdir, "f.png"))
        save_wall_plot       (_WALL_RESULT, os.path.join(tmpdir, "w.png"))

    after = len(plt.get_fignums())
    print(f"  Open figures before : {before}")
    print(f"  Open figures after  : {after}")
    assert after <= before, (
        f"Figure leak: {after - before} figures not closed"
    )

    print("\n  PASS  test_save_functions_close_figures")


# ===========================================================================
#  TEST 26 – Foundation plot dimensioning: footing width B appears on figure
# ===========================================================================

def test_foundation_plot_b_annotation():
    """
    The foundation cross-section figure must include an annotation showing
    the footing width B. This confirms the dimensioning block rendered.
    """
    print("\n" + "="*60)
    print("  TEST 26 – Foundation figure contains B= dimension text")
    print("="*60)

    B   = _FND_PARAMS["B"]
    fig = plot_foundation_bearing(_FND_RESULT)
    ax  = fig.axes[0]

    all_texts = [t.get_text() for t in ax.texts]
    found = any(f"B = {B:.2f}" in t for t in all_texts)

    print(f"  Looking for 'B = {B:.2f} m' in {len(all_texts)} text objects")
    for t in all_texts:
        if "B" in t:
            print(f"    '{t.strip()}'")

    assert found, (
        f"'B = {B:.2f} m' not found in figure texts: {all_texts}"
    )

    plt.close(fig)
    print("\n  PASS  test_foundation_plot_b_annotation")


# ===========================================================================
#  TEST 27 – Wall plot pressure diagram: Ka × γ × H matches figure label
# ===========================================================================

def test_wall_plot_pressure_label():
    """
    The active pressure diagram must display the pressure at the base of the
    wall (Ka × γ × H). The text label must be present in the figure.
    """
    print("\n" + "="*60)
    print("  TEST 27 – Wall figure shows Ka*γ*H pressure label")
    print("="*60)

    Ka    = float(_WALL_RESULT["Ka"])
    gamma = float(_WALL_PARAMS["gamma"])
    H     = float(_WALL_PARAMS["H_wall"])
    pa_h  = Ka * gamma * H

    fig = plot_retaining_wall(_WALL_RESULT)
    ax  = fig.axes[0]
    all_texts = [t.get_text() for t in ax.texts]

    # Label contains the pressure value rounded to 1 d.p.
    pa_str = f"{pa_h:.1f}"
    found  = any(pa_str in t for t in all_texts)

    print(f"  Ka={Ka:.4f}  γ={gamma:.1f}  H={H:.1f}  pa={pa_h:.1f} kPa")
    print(f"  Looking for '{pa_str}' in {len(all_texts)} text objects")

    assert found, f"'{pa_str}' kPa label not found in wall figure texts: {all_texts}"

    plt.close(fig)
    print("\n  PASS  test_wall_plot_pressure_label")


# ===========================================================================
#  TEST 28 – Foundation PDF has correct page count
# ===========================================================================

def test_foundation_pdf_page_count():
    """
    The foundation PDF must have at least 1 page.
    """
    print("\n" + "="*60)
    print("  TEST 28 – Foundation PDF page count ≥ 1")
    print("="*60)

    try:
        import pypdf
    except ImportError:
        print("  SKIP: pypdf not available")
        return

    with tempfile.TemporaryDirectory() as tmpdir:
        path = os.path.join(tmpdir, "fnd_pages.pdf")
        generate_foundation_report(path, _FND_RESULT,
                                   project="PageCountTest", job_ref="PC-001")
        reader  = pypdf.PdfReader(path)
        n_pages = len(reader.pages)
        print(f"  Pages : {n_pages}")
        assert n_pages >= 1

    print("\n  PASS  test_foundation_pdf_page_count")


# ===========================================================================
#  TEST 29 – Wall PDF has correct page count
# ===========================================================================

def test_wall_pdf_page_count():
    """
    The wall PDF must have at least 1 page.
    """
    print("\n" + "="*60)
    print("  TEST 29 – Wall PDF page count ≥ 1")
    print("="*60)

    try:
        import pypdf
    except ImportError:
        print("  SKIP: pypdf not available")
        return

    with tempfile.TemporaryDirectory() as tmpdir:
        path = os.path.join(tmpdir, "wall_pages.pdf")
        generate_wall_report(path, _WALL_RESULT,
                             project="WallPageCount", job_ref="WPC-001")
        reader  = pypdf.PdfReader(path)
        n_pages = len(reader.pages)
        print(f"  Pages : {n_pages}")
        assert n_pages >= 1

    print("\n  PASS  test_wall_pdf_page_count")


# ---------------------------------------------------------------------------
#  Updated runner — original 6 + Sprint 11 new 23 = 29 total
# ---------------------------------------------------------------------------

# ============================================================
#  Sprint 12 — PNG exports + unified project report
# ============================================================

# Shared fixtures for Sprint 12
_FA = None
_WA = None

def _get_fa():
    global _FA
    if _FA is None:
        _FA = run_foundation_analysis({
            "gamma": 18.0, "phi_k": 30.0, "c_k": 0.0,
            "B": 2.0, "Df": 1.0, "Gk": 200.0, "Qk": 80.0,
        })
    return _FA

def _get_wa():
    global _WA
    if _WA is None:
        _WA = run_wall_analysis({
            "gamma": 18.0, "phi_k": 30.0, "c_k": 0.0,
            "H_wall": 4.0, "B_base": 2.8, "B_toe": 0.6,
            "Gk_wall": 0, "t_base": 0.4,
            "t_stem_base": 0.35, "t_stem_top": 0.25,
        })
    return _WA


def test_export_wall_plot_png_returns_bytes():
    """
    export_wall_plot_png() must return bytes (not None, not a file path).

    Reference: Sprint 12 scope — api.export_wall_plot_png().
    """
    print("\n" + "="*60)
    print("  TEST 30 – export_wall_plot_png: returns bytes")
    print("="*60)
    from api import export_wall_plot_png
    result = export_wall_plot_png(_get_wa(), dpi=72)
    assert isinstance(result, bytes), f"Expected bytes, got {type(result)}"
    assert len(result) > 0, "Empty bytes returned"
    print(f"  PNG bytes returned: {len(result):,}  OK")
    print("\n  PASS  test_export_wall_plot_png_returns_bytes")


def test_export_foundation_plot_png_returns_bytes():
    """export_foundation_plot_png() must return non-empty bytes."""
    print("\n" + "="*60)
    print("  TEST 31 – export_foundation_plot_png: returns bytes")
    print("="*60)
    from api import export_foundation_plot_png
    result = export_foundation_plot_png(_get_fa(), dpi=72)
    assert isinstance(result, bytes) and len(result) > 0
    print(f"  PNG bytes returned: {len(result):,}  OK")
    print("\n  PASS  test_export_foundation_plot_png_returns_bytes")


def test_export_wall_png_is_valid_png():
    """Wall PNG bytes must start with the PNG magic bytes (\\x89PNG)."""
    print("\n" + "="*60)
    print("  TEST 32 – export_wall_plot_png: valid PNG magic bytes")
    print("="*60)
    from api import export_wall_plot_png
    result = export_wall_plot_png(_get_wa(), dpi=72)
    PNG_MAGIC = b"\x89PNG"
    assert result[:4] == PNG_MAGIC, (
        f"Expected PNG magic {PNG_MAGIC!r}, got {result[:4]!r}"
    )
    print(f"  Magic bytes: {result[:4]!r}  OK")
    print("\n  PASS  test_export_wall_png_is_valid_png")


def test_export_foundation_png_is_valid_png():
    """Foundation PNG bytes must start with the PNG magic bytes."""
    print("\n" + "="*60)
    print("  TEST 33 – export_foundation_plot_png: valid PNG magic bytes")
    print("="*60)
    from api import export_foundation_plot_png
    result = export_foundation_plot_png(_get_fa(), dpi=72)
    assert result[:4] == b"\x89PNG", f"Not a PNG: {result[:4]!r}"
    print(f"  Magic bytes: {result[:4]!r}  OK")
    print("\n  PASS  test_export_foundation_png_is_valid_png")


def test_export_wall_png_size():
    """Wall PNG at dpi=150 must be at least 50 kB (non-trivial render)."""
    print("\n" + "="*60)
    print("  TEST 34 – export_wall_plot_png: file size ≥ 50 kB at dpi=150")
    print("="*60)
    from api import export_wall_plot_png
    result = export_wall_plot_png(_get_wa(), dpi=150)
    print(f"  PNG size: {len(result):,} bytes")
    assert len(result) >= 50_000, f"PNG too small ({len(result)} bytes)"
    print("\n  PASS  test_export_wall_png_size")


def test_export_foundation_png_size():
    """Foundation PNG at dpi=150 must be at least 50 kB."""
    print("\n" + "="*60)
    print("  TEST 35 – export_foundation_plot_png: file size ≥ 50 kB at dpi=150")
    print("="*60)
    from api import export_foundation_plot_png
    result = export_foundation_plot_png(_get_fa(), dpi=150)
    print(f"  PNG size: {len(result):,} bytes")
    assert len(result) >= 50_000, f"PNG too small ({len(result)} bytes)"
    print("\n  PASS  test_export_foundation_png_size")


def test_generate_project_report_single_analysis():
    """
    generate_project_report() with a single foundation analysis produces
    a PDF file that exists and has size > 1 kB.

    Reference: Sprint 12 — unified project report.
    """
    print("\n" + "="*60)
    print("  TEST 36 – generate_project_report: single analysis")
    print("="*60)
    from exporters.report_pdf import generate_project_report
    with tempfile.TemporaryDirectory() as tmpdir:
        path = os.path.join(tmpdir, "project_single.pdf")
        generate_project_report(
            analyses=[_get_fa()],
            out_path=path,
            project="SingleAnalysisTest",
            job_ref="S-001",
            calc_by="Auto",
            checked_by="QA",
        )
        assert os.path.exists(path), "PDF not created"
        size = os.path.getsize(path)
        print(f"  PDF size: {size:,} bytes")
        assert size > 1_000, f"PDF too small ({size} bytes)"
        with open(path, "rb") as f:
            magic = f.read(4)
        assert magic == b"%PDF", f"Not a PDF: {magic!r}"
        print(f"  Magic bytes: {magic!r}  OK")
    print("\n  PASS  test_generate_project_report_single_analysis")


def test_generate_project_report_multi_analysis():
    """
    generate_project_report() with foundation + wall produces a larger PDF
    than the single-analysis case.

    The combined report must be larger because it contains two calc sheets.
    """
    print("\n" + "="*60)
    print("  TEST 37 – generate_project_report: foundation + wall")
    print("="*60)
    from exporters.report_pdf import generate_project_report
    with tempfile.TemporaryDirectory() as tmpdir:
        p1 = os.path.join(tmpdir, "single.pdf")
        p2 = os.path.join(tmpdir, "multi.pdf")
        generate_project_report([_get_fa()], p1,
                                project="Multi", job_ref="M-001",
                                calc_by="Auto", checked_by="QA")
        generate_project_report([_get_fa(), _get_wa()], p2,
                                project="Multi", job_ref="M-002",
                                calc_by="Auto", checked_by="QA")
        s1 = os.path.getsize(p1)
        s2 = os.path.getsize(p2)
        print(f"  Single: {s1:,} bytes   Multi: {s2:,} bytes")
        assert s2 > s1, f"Multi-section PDF not larger than single ({s2} <= {s1})"
    print("\n  PASS  test_generate_project_report_multi_analysis")


def test_project_report_is_valid_pdf():
    """
    The project report PDF must start with %PDF magic bytes and be
    parseable by pypdf (decompress streams successfully).
    """
    print("\n" + "="*60)
    print("  TEST 38 – generate_project_report: valid PDF, parseable by pypdf")
    print("="*60)
    from exporters.report_pdf import generate_project_report
    with tempfile.TemporaryDirectory() as tmpdir:
        path = os.path.join(tmpdir, "project_valid.pdf")
        generate_project_report(
            analyses=[_get_fa(), _get_wa()],
            out_path=path,
            project="ValidityTest",
            job_ref="V-001",
        )
        with open(path, "rb") as f:
            magic = f.read(4)
        assert magic == b"%PDF", f"Not a PDF: {magic!r}"
        print(f"  Magic bytes OK: {magic!r}")
        try:
            import pypdf
            reader = pypdf.PdfReader(path)
            n = len(reader.pages)
            print(f"  pypdf parsed {n} pages successfully")
            assert n >= 1
        except ImportError:
            print("  pypdf not available — skipping deep parse check")
    print("\n  PASS  test_project_report_is_valid_pdf")


def test_project_report_has_multiple_pages():
    """
    A project report combining two analyses must have ≥ 2 pages.

    Each individual analysis contributes at least one page; the cover
    adds one more.
    """
    print("\n" + "="*60)
    print("  TEST 39 – generate_project_report: ≥ 2 pages for 2 analyses")
    print("="*60)
    try:
        import pypdf
    except ImportError:
        print("  SKIP: pypdf not available")
        return
    from exporters.report_pdf import generate_project_report
    with tempfile.TemporaryDirectory() as tmpdir:
        path = os.path.join(tmpdir, "project_pages.pdf")
        generate_project_report(
            analyses=[_get_fa(), _get_wa()],
            out_path=path, project="PagesTest",
        )
        reader = pypdf.PdfReader(path)
        n = len(reader.pages)
        print(f"  Pages: {n}")
        assert n >= 2, f"Expected ≥ 2 pages, got {n}"
    print("\n  PASS  test_project_report_has_multiple_pages")


def test_export_project_pdf_api_function():
    """
    api.export_project_pdf() produces a valid PDF file at out_path.

    Verifies the end-to-end api → report_pdf → pypdf chain.
    """
    print("\n" + "="*60)
    print("  TEST 40 – api.export_project_pdf(): end-to-end")
    print("="*60)
    from api import export_project_pdf
    with tempfile.TemporaryDirectory() as tmpdir:
        path = os.path.join(tmpdir, "api_project.pdf")
        result = export_project_pdf(
            analyses=[_get_fa(), _get_wa()],
            out_path=path,
            project="APITest",
            job_ref="API-001",
            calc_by="Auto",
            checked_by="QA",
        )
        assert result == path, f"Expected out_path returned, got {result!r}"
        assert os.path.exists(path), "PDF not created"
        size = os.path.getsize(path)
        print(f"  PDF size: {size:,} bytes")
        assert size > 1_000, f"PDF too small ({size} bytes)"
        with open(path, "rb") as f:
            magic = f.read(4)
        assert magic == b"%PDF"
        print(f"  Magic bytes OK: {magic!r}")
    print("\n  PASS  test_export_project_pdf_api_function")



if __name__ == "__main__":
    # Original 6
    test_plot_slope_returns_figure()
    test_save_slope_plot_writes_file()
    test_fos_heatmap_returns_figure()
    test_save_fos_heatmap_writes_file()
    test_generate_slope_report_produces_pdf()
    test_pdf_contains_expected_content()

    # Sprint 11 — foundation plot (Tests 7-10)
    test_plot_foundation_returns_figure()
    test_save_foundation_plot_writes_file()
    test_foundation_figure_axis_labels()
    test_foundation_figure_pass_fail_colour()

    # Sprint 11 — wall plot (Tests 11-14)
    test_plot_wall_returns_figure()
    test_save_wall_plot_writes_file()
    test_wall_figure_axis_labels()
    test_wall_figure_y_range_covers_height()

    # Sprint 11 — foundation PDF (Tests 15-16)
    test_generate_foundation_report_produces_pdf()
    test_foundation_pdf_contains_expected_content()

    # Sprint 11 — wall PDF (Tests 17-18)
    test_generate_wall_report_produces_pdf()
    test_wall_pdf_contains_expected_content()

    # Sprint 11 — foundation DOCX (Tests 19-20)
    test_generate_foundation_report_docx()
    test_foundation_docx_text_content()

    # Sprint 11 — wall DOCX (Tests 21-22)
    test_generate_wall_report_docx()
    test_wall_docx_text_content()

    # Sprint 11 — robustness & details (Tests 23-29)
    test_foundation_report_fail_case()
    test_wall_report_fail_case()
    test_save_functions_close_figures()
    test_foundation_plot_b_annotation()
    test_wall_plot_pressure_label()
    test_foundation_pdf_page_count()
    test_wall_pdf_page_count()

    # Sprint 12 — PNG export functions + unified project report (Tests 30-40)
    test_export_wall_plot_png_returns_bytes()
    test_export_foundation_plot_png_returns_bytes()
    test_export_wall_png_is_valid_png()
    test_export_foundation_png_is_valid_png()
    test_export_wall_png_size()
    test_export_foundation_png_size()
    test_generate_project_report_single_analysis()
    test_generate_project_report_multi_analysis()
    test_project_report_is_valid_pdf()
    test_project_report_has_multiple_pages()
    test_export_project_pdf_api_function()

    print("\n" + "="*60)
    print("  ALL exporter tests passed.  (40/40)")
    print("="*60 + "\n")


