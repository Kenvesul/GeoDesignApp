"""
tests/test_report_docx.py
=========================
Validation tests for exporters/report_docx.py.

Tests
-----
1. generate_slope_report_docx() produces a valid .docx file (non-empty, ZIP magic bytes)
2. File size is substantial (> 30 kB — ensures figure and tables are rendered)
3. All required text tokens appear in document body (via python-docx extraction)
4. Document contains at least the expected number of tables (header + inputs + DA1 + slices)
5. Document contains an embedded image (cross-section figure)
6. DA1 verdict text is present (SATISFACTORY or UNSATISFACTORY)
"""

import sys, os, math, tempfile
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from models.soil import Soil
from models.geometry import SlopeGeometry
from core.search import grid_search
from core.factors_of_safety import verify_slope_da1
from core.slicer import create_slices
from exporters.report_docx import generate_slope_report_docx

# ── Shared fixture ────────────────────────────────────────────────────────────
#  1:2 slope with dense sand — matches test_exporters.py fixture for consistency
SLOPE = SlopeGeometry([(0, 3), (6, 3), (12, 0), (18, 0)])
SOIL  = Soil("Dense Sand", 19.0, 35.0, 0.0)
RU    = 0.0

_RESULT = grid_search(
    SLOPE, SOIL, ru=RU,
    cx_range=(3.0, 14.0), cy_range=(3.0, 12.0), r_range=(4.0, 16.0),
    n_cx=8, n_cy=8, n_r=6,
)
_VERIFICATION = verify_slope_da1(SLOPE, SOIL, ru=RU)
_SLICES       = create_slices(SLOPE, _RESULT.critical_circle, SOIL, num_slices=20)


def _generate(path: str, project: str = "DocxTest", job_ref: str = "DT-001") -> None:
    generate_slope_report_docx(
        filepath      = path,
        soil          = SOIL,
        slope         = SLOPE,
        search_result = _RESULT,
        verification  = _VERIFICATION,
        slices        = _SLICES,
        ru            = RU,
        project       = project,
        job_ref       = job_ref,
        calc_by       = "Auto",
        checked_by    = "QA",
    )


# ── Test 1: valid .docx produced ─────────────────────────────────────────────
def test_docx_file_is_valid_zip():
    """
    A .docx file is a ZIP archive whose first two bytes are PK (0x50 0x4B).
    python-docx raises BadZipFile if the archive is corrupt.
    """
    print("\n" + "="*60)
    print("  TEST 1 – generate_slope_report_docx: valid .docx produced")
    print("="*60)

    with tempfile.TemporaryDirectory() as tmp:
        path = os.path.join(tmp, "slope.docx")
        _generate(path)

        raw = open(path, "rb").read()
        assert raw[:2] == b"PK", \
            f"FAIL: first bytes are {raw[:2]!r}, expected b'PK' (ZIP/DOCX magic)"
        print(f"  ZIP magic bytes : OK ({raw[:2]})")

        # Verify python-docx can re-open it
        from docx import Document as _Doc
        doc = _Doc(path)
        assert len(doc.paragraphs) > 0, "FAIL: document has no paragraphs"
        print(f"  Re-opened OK   : {len(doc.paragraphs)} paragraphs")

    print("\n  PASS  test_docx_file_is_valid_zip")


# ── Test 2: file size is substantial ─────────────────────────────────────────
def test_docx_file_size():
    """
    File must be > 30 kB, ensuring tables and the embedded PNG figure
    are present (blank/empty documents are typically < 10 kB).
    """
    print("\n" + "="*60)
    print("  TEST 2 – file size > 30 kB (non-trivial render)")
    print("="*60)

    with tempfile.TemporaryDirectory() as tmp:
        path = os.path.join(tmp, "slope.docx")
        _generate(path)
        size = os.path.getsize(path)
        print(f"  File size : {size:,} bytes")
        assert size > 30_000, \
            f"FAIL: file too small ({size} bytes) — figure or tables likely missing"

    print("\n  PASS  test_docx_file_size")


# ── Test 3: required text tokens present ─────────────────────────────────────
def test_docx_contains_expected_tokens():
    """
    Extract all paragraph text and verify key tokens exist:
      - project name, job ref, soil name, 'DA1', 'EC7'
    """
    print("\n" + "="*60)
    print("  TEST 3 – required text tokens found in document body")
    print("="*60)

    from docx import Document as _Doc

    PROJECT = "TokenCheckProject"
    JOB_REF = "TCR-999"

    with tempfile.TemporaryDirectory() as tmp:
        path = os.path.join(tmp, "slope.docx")
        _generate(path, project=PROJECT, job_ref=JOB_REF)

        doc  = _Doc(path)
        # Gather text from paragraphs AND table cells
        texts = []
        for para in doc.paragraphs:
            texts.append(para.text)
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        texts.append(para.text)
        full_text = "\n".join(texts)

        tokens = [PROJECT, JOB_REF, "Dense Sand", "DA1", "EC7"]
        for tok in tokens:
            found = tok in full_text
            print(f"  Token '{tok}' : {'FOUND' if found else 'MISSING'}")
            assert found, (
                f"FAIL: token '{tok}' not found in document text.\n"
                f"  First 500 chars:\n{full_text[:500]}"
            )

    print("\n  PASS  test_docx_contains_expected_tokens")


# ── Test 4: expected number of tables ────────────────────────────────────────
def test_docx_table_count():
    """
    Document must contain at least 4 tables:
      header block, input parameters, DA1 verification, slice summary.
    """
    print("\n" + "="*60)
    print("  TEST 4 – document contains ≥ 4 tables")
    print("="*60)

    from docx import Document as _Doc

    with tempfile.TemporaryDirectory() as tmp:
        path = os.path.join(tmp, "slope.docx")
        _generate(path)
        doc    = _Doc(path)
        n_tbls = len(doc.tables)
        print(f"  Table count : {n_tbls}")
        assert n_tbls >= 4, \
            f"FAIL: only {n_tbls} tables — expected ≥ 4 (header, inputs, DA1, slices)"

    print("\n  PASS  test_docx_table_count")


# ── Test 5: embedded image present ───────────────────────────────────────────
def test_docx_contains_image():
    """
    The cross-section figure must be embedded in the document's media store.
    Checked by inspecting the ZIP archive for image entries.
    """
    print("\n" + "="*60)
    print("  TEST 5 – embedded cross-section image present")
    print("="*60)

    import zipfile

    with tempfile.TemporaryDirectory() as tmp:
        path = os.path.join(tmp, "slope.docx")
        _generate(path)

        with zipfile.ZipFile(path) as zf:
            media_files = [n for n in zf.namelist() if n.startswith("word/media/")]

        print(f"  Media entries : {media_files}")
        assert len(media_files) >= 1, \
            "FAIL: no media files in docx — embedded PNG figure is missing"

    print("\n  PASS  test_docx_contains_image")


# ── Test 6: DA1 verdict text present ─────────────────────────────────────────
def test_docx_verdict_text():
    """
    The document must contain either 'SATISFACTORY' or 'UNSATISFACTORY'
    in the verdict paragraph — confirming EC7 DA1 logic is wired up.
    """
    print("\n" + "="*60)
    print("  TEST 6 – DA1 verdict text ('SATISFACTORY' / 'UNSATISFACTORY')")
    print("="*60)

    from docx import Document as _Doc

    with tempfile.TemporaryDirectory() as tmp:
        path = os.path.join(tmp, "slope.docx")
        _generate(path)

        doc  = _Doc(path)
        all_text = "\n".join(p.text for p in doc.paragraphs)
        found = ("SATISFACTORY" in all_text) or ("UNSATISFACTORY" in all_text)
        expected = "SATISFACTORY" if _VERIFICATION.passes else "UNSATISFACTORY"
        present  = expected in all_text
        print(f"  Expected verdict : {expected}")
        print(f"  Found in text    : {present}")
        assert present, \
            f"FAIL: '{expected}' not found in document paragraphs.\n{all_text[:400]}"

    print("\n  PASS  test_docx_verdict_text")


# ── Runner ────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    tests = [
        test_docx_file_is_valid_zip,
        test_docx_file_size,
        test_docx_contains_expected_tokens,
        test_docx_table_count,
        test_docx_contains_image,
        test_docx_verdict_text,
    ]
    passed = failed = 0
    for fn in tests:
        try:
            fn()
            passed += 1
        except AssertionError as exc:
            print(f"\n  *** FAIL ***  {fn.__name__}")
            print(f"  {exc}")
            failed += 1
        except Exception as exc:
            print(f"\n  *** ERROR ***  {fn.__name__}: {exc}")
            import traceback; traceback.print_exc()
            failed += 1

    print("\n" + "="*60)
    if failed == 0:
        print(f"  ALL {passed} docx tests PASSED.")
    else:
        print(f"  {failed} FAILED / {passed} passed.")
    print("="*60)
    sys.exit(failed)
